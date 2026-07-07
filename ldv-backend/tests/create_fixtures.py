"""
Generate all test fixtures for the Sydeco LightML Contract Risk Analyzer validation pack.
Run once from ldv-backend/: python tests/create_fixtures.py
"""
import struct
from pathlib import Path

import fitz
import docx

FIXTURES = Path(__file__).parent / "fixtures"
for sub in ("pdf", "docx", "txt", "negative"):
    (FIXTURES / sub).mkdir(parents=True, exist_ok=True)


# ── helpers ───────────────────────────────────────────────────────────────────

def make_pdf(path: Path, title: str, body: str) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), f"{title}\n\n{body}", fontsize=10)
    doc.save(str(path))
    doc.close()


def make_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    d = docx.Document()
    d.add_heading(title, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(str(path))


def make_txt(path: Path, content: str, encoding: str = "utf-8") -> None:
    path.write_bytes(content.encode(encoding))


# ── PDF fixtures ──────────────────────────────────────────────────────────────

make_pdf(
    FIXTURES / "pdf" / "01_employment_id.pdf",
    "PERJANJIAN KERJA",
    """Antara PT Maju Bersama (Pemberi Kerja) dan Budi Santoso (Pekerja).

Pasal 1 - Identitas Para Pihak
Pemberi Kerja: PT Maju Bersama, Jakarta. Pekerja: Budi Santoso, 1 Jan 1990.

Pasal 2 - Jabatan dan Uraian Kerja
Pekerja diangkat sebagai Software Engineer sesuai UU Ketenagakerjaan No. 13 Tahun 2003.

Pasal 3 - Masa Kerja
Kontrak berlaku 1 April 2026 sampai 31 Maret 2027.

Pasal 4 - Gaji dan Tunjangan
Gaji pokok Rp 12.000.000/bulan, dibayar tanggal 25 setiap bulan.

Pasal 5 - Jam Kerja
40 jam/minggu sesuai Peraturan Menteri Ketenagakerjaan.

Pasal 6 - Cuti dan Hari Libur
12 hari cuti tahunan sesuai ketentuan UU.

Pasal 7 - Pemutusan Hubungan Kerja
PHK mengikuti ketentuan UU Ketenagakerjaan dan wajib memberikan pesangon.

Pasal 8 - Kerahasiaan
Pekerja wajib menjaga kerahasiaan informasi perusahaan.

Ditandatangani di Jakarta, 1 April 2026.
Pemberi Kerja: ________________    Pekerja: ________________""",
)

make_pdf(
    FIXTURES / "pdf" / "02_lease_be.pdf",
    "BAIL COMMERCIAL",
    """Entre Immobilière Bruxelles SA (Bailleur) et Jean Dupont (Locataire).

Article 1 - Identité des Parties
Bailleur: Immobilière Bruxelles SA, Belgique. Locataire: Jean Dupont, né 5 mai 1985.

Article 2 - Objet du Contrat
Location d'un local commercial sis rue de la Loi 42, Bruxelles, selon le Code civil Belgique.

Article 3 - Durée du Bail
Le bail prend effet le 1er avril 2026 pour une durée de 3 ans, conformément à la loi belge.

Article 4 - Loyer
Loyer mensuel de 1 800 EUR, payable le 5 de chaque mois.

Article 5 - Paiement
Virement bancaire uniquement. Dépôt de garantie: 2 mois de loyer.

Article 6 - État des Lieux
Un état des lieux contradictoire sera dressé à l'entrée et à la sortie du locataire.

Article 7 - Résiliation Unilatérale
Chaque partie peut résilier avec un préavis de 3 mois selon la loi.

Article 8 - Indexation du Loyer
Le loyer est indexé annuellement selon l'indice santé belge.

Signé à Bruxelles, le 1er avril 2026.
Bailleur: ________________    Locataire: ________________""",
)

make_pdf(
    FIXTURES / "pdf" / "03_nda_en.pdf",
    "NON-DISCLOSURE AGREEMENT",
    """Between Alpha Corp (Disclosing Party) and Beta Ltd (Receiving Party).

1. Purpose
The parties intend to explore a potential business collaboration. Beta Ltd may receive
confidential information from Alpha Corp for evaluation purposes only.

2. Confidential Information
All technical, financial, and business data shared under this agreement is confidential.

3. Obligations
The Receiving Party shall not disclose, copy, or use confidential information outside
the stated purpose without written consent.

4. Duration
This agreement is effective for 2 years from the date of signing.

5. Termination
Either party may terminate this agreement with 30 days' written notice.

6. Governing Law
This agreement shall be governed by the laws of England and Wales.

Signed on 1 April 2026.
Alpha Corp: ________________    Beta Ltd: ________________""",
)

make_pdf(
    FIXTURES / "pdf" / "04_incomplete_en.pdf",
    "SERVICE CONTRACT (DRAFT)",
    """Between TechStudio Inc and Client ABC.

1. Services
TechStudio Inc agrees to deliver a custom software application.

2. Fees
The client agrees to pay USD 50,000 upon delivery.

[NOTE: Signatures section missing]""",
)

make_pdf(
    FIXTURES / "pdf" / "05_brochure_en.pdf",
    "PRODUCT BROCHURE – CloudSuite 2026",
    """Discover the future of enterprise productivity with CloudSuite 2026.

Key Features:
- Real-time collaboration across 50+ apps
- AI-powered workflow automation
- 99.99% uptime SLA guarantee
- ISO 27001 certified security

Pricing Plans:
Starter: USD 29/month per user
Professional: USD 79/month per user
Enterprise: Contact our sales team

CloudSuite is used by over 10,000 businesses worldwide.
Visit www.cloudsuite.example.com or call +1-800-CLOUD-26.

This brochure is for informational purposes only and does not constitute a contract.""",
)

# ponytail: blank page, no insert_text call — simulates a scanned image with
# no OCR text layer, exercising app.py's existing "Scan/OCR required" path.
_scan_doc = fitz.open()
_scan_doc.new_page()
_scan_doc.save(str(FIXTURES / "pdf" / "06_scanned_blank_en.pdf"))
_scan_doc.close()


# ── DOCX fixtures ─────────────────────────────────────────────────────────────

make_docx(
    FIXTURES / "docx" / "01_service_agreement_en.docx",
    "SERVICE AGREEMENT",
    [
        "Between DataFlow Solutions Ltd (Service Provider) and RetailCo GmbH (Client).",
        "1. Identity of the Parties\nService Provider: DataFlow Solutions Ltd. Client: RetailCo GmbH.",
        "2. Subject Matter\nService Provider shall deliver data integration services as described in Annex A.",
        "3. Price and Payment Terms\nTotal fee: EUR 24,000 payable in quarterly instalments of EUR 6,000.",
        "4. Delivery and Performance\nDelivery within 90 days of signing. Milestones as per Annex B.",
        "5. Warranties\nService Provider warrants that services will meet the specifications in Annex A.",
        "6. Liability Limitation\nNeither party shall be liable for indirect or consequential damages.",
        "7. Dispute Resolution\nDisputes shall be resolved by arbitration under ICC Rules in Paris.",
        "8. Governing Law\nThis agreement is governed by the laws of France.",
        "Signed on 1 April 2026.\nService Provider: ________________    Client: ________________",
    ],
)

make_docx(
    FIXTURES / "docx" / "02_employment_fr.docx",
    "CONTRAT DE TRAVAIL",
    [
        "Entre Agence Conseil Paris SAS (Employeur) et Sophie Martin (Employée).",
        "Article 1 - Identité des parties\nEmployeur: Agence Conseil Paris SAS. Employée: Sophie Martin.",
        "Article 2 - Poste et description\nSophie Martin est engagée en qualité de Chef de Projet selon le Code du travail France.",
        "Article 3 - Date de début et durée\nContrat à durée indéterminée, prenant effet le 1er avril 2026.",
        "Article 4 - Salaire\nSalaire brut mensuel de 4 200 EUR, versé le dernier jour ouvrable du mois.",
        "Article 5 - Temps de travail\n35 heures par semaine conformément à la loi française.",
        "Article 6 - Congés\n25 jours de congés payés par an selon la loi.",
        "Article 7 - Rupture du contrat\nPréavis de 3 mois pour chaque partie selon le Code du travail.",
        "Article 8 - Confidentialité\nL'employée s'engage à respecter la confidentialité des informations de l'entreprise.",
        "Signé à Paris, le 1er avril 2026.\nEmployeur: ________________    Employée: ________________",
    ],
)

make_docx(
    FIXTURES / "docx" / "03_nda_nl.docx",
    "GEHEIMHOUDINGSOVEREENKOMST",
    [
        "Tussen Innovatie BV (Openbarende Partij) en Consultancy Nederland BV (Ontvangende Partij).",
        "Artikel 1 - Identiteit van de Partijen\nInnovatie BV, gevestigd in Amsterdam, Nederland. Consultancy Nederland BV, gevestigd in Rotterdam.",
        "Artikel 2 - Doel\nPartijen willen samenwerken op het gebied van productontwikkeling. Dit BW-contract regelt de geheimhouding.",
        "Artikel 3 - Vertrouwelijke Informatie\nAlle technische en commerciële informatie die wordt gedeeld valt onder deze overeenkomst.",
        "Artikel 4 - Verplichtingen\nDe Ontvangende Partij zal geen vertrouwelijke informatie openbaar maken zonder schriftelijke toestemming.",
        "Artikel 5 - Duur\nDeze overeenkomst geldt voor een periode van 3 jaar.",
        "Artikel 6 - Toepasselijk recht\nDeze overeenkomst wordt beheerst door het Nederlands recht (wet BW).",
        "Ondertekend te Amsterdam, 1 april 2026.\nInnovatie BV: ________________    Consultancy Nederland BV: ________________",
    ],
)

make_docx(
    FIXTURES / "docx" / "04_legal_memo_en.docx",
    "LEGAL MEMORANDUM",
    [
        "TO: Board of Directors, Nexus Holdings",
        "FROM: Legal Department",
        "RE: Compliance with Data Protection Regulations",
        "DATE: 1 April 2026",
        "This memorandum summarises the obligations of Nexus Holdings under applicable data "
        "protection law and recommends immediate corrective actions.",
        "1. Current Status\nThe company currently lacks a formal data retention policy and has not "
        "appointed a Data Protection Officer as required by applicable regulations.",
        "2. Recommended Actions\n(a) Appoint a DPO within 30 days.\n(b) Draft and publish a data retention policy.\n"
        "(c) Conduct a full audit of data processing activities.",
        "This memo is privileged and confidential.",
    ],
)

make_docx(
    FIXTURES / "docx" / "05_general_terms_en.docx",
    "GENERAL TERMS AND CONDITIONS",
    [
        "These General Terms and Conditions apply to all services provided by OmniServices Ltd.",
        "1. Definitions\n'Service' means any professional service delivered by OmniServices Ltd to a Client.",
        "2. Orders\nAll orders must be confirmed in writing before work commences.",
        "3. Payment\nInvoices are due within 30 days of issue. Late payments accrue interest at 8% per annum.",
        "4. Intellectual Property\nAll deliverables remain the property of OmniServices Ltd until full payment is received.",
        "5. Limitation of Liability\nOmniServices Ltd's total liability shall not exceed the fees paid in the preceding 3 months.",
        "6. Governing Law\nThese terms are governed by English law.",
    ],
)

make_docx(
    FIXTURES / "docx" / "06_memo_fr.docx",
    "NOTE JURIDIQUE INTERNE",
    [
        "À : Conseil d'Administration, Société Lumière SAS",
        "De : Département Juridique",
        "Objet : Conformité au Règlement Général sur la Protection des Données (RGPD)",
        "Date : 1 avril 2026",
        "Cette note résume les obligations de Société Lumière SAS en vertu du RGPD et "
        "recommande des mesures correctives immédiates.",
        "1. Situation actuelle\nLa société ne dispose actuellement d'aucune politique formelle "
        "de conservation des données et n'a pas désigné de délégué à la protection des données.",
        "2. Actions recommandées\n(a) Désigner un DPO dans les 30 jours.\n"
        "(b) Rédiger et publier une politique de conservation des données.\n"
        "(c) Mener un audit complet des activités de traitement des données.",
        "Cette note est confidentielle et protégée par le secret professionnel.",
    ],
)

make_docx(
    FIXTURES / "docx" / "07_brochure_nl.docx",
    "PRODUCTBROCHURE – CloudSuite 2026",
    [
        "Ontdek de toekomst van productiviteit voor bedrijven met CloudSuite 2026.",
        "Belangrijkste functies:\n- Realtime samenwerking in meer dan 50 apps\n"
        "- AI-gestuurde workflow-automatisering\n- 99,99% uptime garantie\n"
        "- ISO 27001 gecertificeerde beveiliging",
        "Prijsplannen:\nStarter: EUR 29/maand per gebruiker\nProfessional: EUR 79/maand per gebruiker\n"
        "Enterprise: Neem contact op met onze verkoopafdeling",
        "CloudSuite wordt gebruikt door meer dan 10.000 bedrijven wereldwijd.",
    ],
)


# ── TXT fixtures ──────────────────────────────────────────────────────────────

make_txt(
    FIXTURES / "txt" / "01_employment_id.txt",
    """PERJANJIAN KERJA

Antara CV Teknologi Nusantara (Pemberi Kerja) dan Dewi Rahayu (Pekerja).

Pasal 1 - Identitas Para Pihak
CV Teknologi Nusantara beralamat di Surabaya. Pekerja: Dewi Rahayu.

Pasal 2 - Jabatan
Pekerja diangkat sebagai Data Analyst berdasarkan UU Ketenagakerjaan No. 13 Tahun 2003.

Pasal 3 - Masa Kerja
1 Mei 2026 hingga 30 April 2027.

Pasal 4 - Gaji
Rp 9.000.000 per bulan.

Pasal 5 - Jam Kerja
Senin-Jumat, 08.00-17.00, sesuai Peraturan Menteri Ketenagakerjaan.

Pasal 6 - Cuti
12 hari kerja per tahun.

Pasal 7 - PHK
Sesuai ketentuan UU yang berlaku.

Pasal 8 - Tanda Tangan
Pemberi Kerja: ________________    Pekerja: ________________
""",
)

make_txt(
    FIXTURES / "txt" / "02_lease_be.txt",
    """BAIL D'HABITATION

Entre M. Pierre Leblanc (Bailleur) et Mme Claire Fontaine (Locataire).

Article 1 - Identité
Bailleur: Pierre Leblanc, Belgique. Locataire: Claire Fontaine.

Article 2 - Objet
Location d'un appartement de 2 chambres, avenue Louise 88, Bruxelles, loi belge.

Article 3 - Durée
Bail de 3 ans à partir du 1er mai 2026, Code civil Belgique.

Article 4 - Loyer
1 200 EUR/mois, payable le 1er de chaque mois.

Article 5 - Dépôt de garantie
2 mois de loyer, soit 2 400 EUR.

Article 6 - Indexation
Selon l'indice santé belge, employé annuellement.

Article 7 - Résiliation
Préavis de 3 mois pour le locataire, 6 mois pour le bailleur.

Signatures:
Bailleur: ________________    Locataire: ________________
""",
    encoding="latin-1",
)

make_txt(
    FIXTURES / "txt" / "03_short_contract_en.txt",
    """FREELANCE CONTRACT

Between DesignStudio Ltd (Client) and Mark Hughes (Freelancer).

1. Services: Logo design and brand identity package.
2. Fee: GBP 2,500 payable 50% upfront and 50% on delivery.
3. Timeline: 4 weeks from signing.
4. Ownership: All designs transfer to Client upon full payment.
5. Governing law: Laws of England and Wales.

Signed: ________________    Date: 1 April 2026
""",
)

make_txt(
    FIXTURES / "txt" / "04_long_agreement_en.txt",
    """MASTER SERVICES AGREEMENT

This Master Services Agreement ("Agreement") is entered into as of 1 April 2026,
between Global Tech Partners Inc, a Delaware corporation ("Company"), and Enterprise
Solutions GmbH, a German company ("Client").

1. DEFINITIONS
"Services" means the professional IT consulting and development services described
in one or more Statements of Work ("SOW") executed under this Agreement.
"Confidential Information" means any non-public information disclosed by either party.

2. SERVICES
Company will perform the Services as described in each SOW. Company may subcontract
work with prior written approval from Client.

3. COMPENSATION
Client shall pay Company according to the rates specified in each SOW.
Invoices are due within 30 days. Late payments accrue 1.5% monthly interest.

4. INTELLECTUAL PROPERTY
Work product created specifically for Client under each SOW shall be assigned to
Client upon full payment. Company retains rights to all pre-existing materials.

5. CONFIDENTIALITY
Each party agrees to protect the other's Confidential Information with the same
degree of care it uses for its own, but no less than reasonable care.

6. WARRANTIES
Company warrants that Services will be performed in a professional manner consistent
with industry standards. Client warrants it has the right to provide all materials.

7. LIMITATION OF LIABILITY
Neither party shall be liable for indirect, special, or consequential damages.
Company's total liability shall not exceed fees paid in the 3 months preceding the claim.

8. TERM AND TERMINATION
This Agreement commences on the Effective Date and continues for 2 years unless
terminated earlier. Either party may terminate with 30 days written notice.
Company may terminate immediately for non-payment after 15-day cure period.

9. DISPUTE RESOLUTION
Disputes shall first be subject to good-faith negotiation for 30 days, then binding
arbitration under AAA rules in New York, New York.

10. GOVERNING LAW
This Agreement is governed by the laws of the State of New York.

11. ENTIRE AGREEMENT
This Agreement constitutes the entire agreement between the parties and supersedes
all prior discussions and agreements.

IN WITNESS WHEREOF the parties have executed this Agreement as of the date first written.

Global Tech Partners Inc: ________________    Enterprise Solutions GmbH: ________________
""",
)

make_txt(
    FIXTURES / "txt" / "05_irrelevant_en.txt",
    """QUARTERLY EARNINGS REPORT – Q1 2026

Revenue grew 12% year-over-year to USD 4.3 billion in Q1 2026.
Operating margins improved to 18.4%, up from 16.1% in the same period last year.

Key highlights:
- Cloud segment revenue: USD 1.8 billion (+34% YoY)
- Hardware segment: USD 900 million (-5% YoY)
- Services segment: USD 1.6 billion (+8% YoY)

The board approved a share buyback programme of USD 500 million.
Dividend declared: USD 0.42 per share, payable 15 May 2026.

Guidance for Q2 2026: revenue expected between USD 4.5 and 4.7 billion.
""",
)

# Additional fixtures for risk scorer bootstrap diversity
make_txt(
    FIXTURES / "txt" / "06_high_risk_leonine_en.txt",
    """SERVICE CONTRACT

Between MegaCorp Inc (Provider) and SmallBiz Ltd (Client).

1. Services
Provider will deliver marketing services at its sole discretion.

2. Payment
Client shall pay USD 10,000 per month.
A penalty of 25% per day applies to late payments.
Late payment fee of 30% per month compounded daily on all outstanding balances.

3. Rights Waiver
Client irrevocably waives all rights to dispute any invoice or charge issued by Provider.
Client surrenders all legal rights under applicable law upon signing this agreement.

4. Liability
Provider accepts no liability whatsoever for any damages, losses or claims of any kind.

5. Modification
Provider may modify the terms of this contract at any time without notice to Client.

6. Profits
All profits shall be allocated exclusively to Provider regardless of Client's contribution.

Signed:
""",
)

make_txt(
    FIXTURES / "txt" / "07_high_risk_missing_clauses_en.txt",
    """CONSULTING AGREEMENT

Between Strategy Partners LLC and Apex Holdings.

1. Services
[Description of services to be confirmed at a later stage.]

2. Payment
USD [AMOUNT TBD] payable as per separate schedule to be agreed.

3. Term
[Insert start and end date here.]

4. Governing Law
[Governing law clause to be added by legal team.]

5. Dispute Resolution
[Dispute resolution mechanism to be negotiated between the parties.]

6. Termination
[Termination clause pending review by counsel.]

7. Intellectual Property
Section 7 — Intellectual Property: [DRAFT — INCOMPLETE]

Signatures: [TO BE COMPLETED]
""",
)

make_txt(
    FIXTURES / "txt" / "08_medium_risk_partial_en.txt",
    """SOFTWARE LICENSE AGREEMENT

Between SoftTech Ltd (Licensor) and RetailGroup Inc (Licensee).

1. License Grant
Licensor grants Licensee a non-exclusive, non-transferable license to use the Software.

2. Payment
Annual license fee of USD 12,000, payable within 30 days of invoice.
Late payments accrue interest at 18% per annum.

3. Restrictions
Licensee shall not sublicense, copy, or reverse engineer the Software.

4. Warranties
Licensor warrants that the Software will substantially perform as described.

5. Governing Law
This agreement shall be governed by the laws of the State of California.

Signed 1 April 2026.
""",
)

make_txt(
    FIXTURES / "txt" / "09_medium_risk_no_venue_en.txt",
    """LOAN AGREEMENT

Between First Capital Bank (Lender) and Sunrise Retail Ltd (Borrower).

1. Loan Amount
Lender agrees to lend USD 250,000 to Borrower.

2. Interest Rate
Interest at 8% per annum, calculated monthly.

3. Repayment
Borrower shall repay the loan in 36 equal monthly instalments beginning 1 June 2026.

4. Security
The loan is secured by a first charge over Borrower's inventory.

5. Default
In the event of default, the full outstanding balance becomes immediately due.

6. Governing Law
This agreement is governed by the laws of New York State.

Signed 1 April 2026.
""",
)

make_txt(
    FIXTURES / "txt" / "10_low_risk_complete_en.txt",
    """MASTER SERVICES AGREEMENT

Between Precision Consulting Group Ltd (Consultant) and BuildCo Holdings Ltd (Client).

1. Parties
Consultant: Precision Consulting Group Ltd, London, England.
Client: BuildCo Holdings Ltd, Manchester, England.

2. Services
Consultant will provide project management consulting as detailed in Schedule A.

3. Fees and Payment
Fees are as set out in Schedule A. Invoices are payable within 30 days.
Late payments accrue interest at 4% per annum above the Bank of England base rate.

4. Intellectual Property
All deliverables produced specifically for Client become Client's property upon full payment.

5. Confidentiality
Both parties agree to keep all information exchanged under this agreement strictly confidential.

6. Limitation of Liability
Neither party is liable for indirect or consequential loss. Total liability is capped at fees paid.

7. Termination
Either party may terminate with 30 days written notice. Immediate termination for material breach.

8. Dispute Resolution
Disputes shall be resolved by mediation, then arbitration under LCIA rules in London.

9. Governing Law
This agreement is governed by the laws of England and Wales.

Signed: 1 April 2026.
""",
)

make_txt(
    FIXTURES / "txt" / "11_low_risk_employment_en.txt",
    """EMPLOYMENT CONTRACT

Between Horizon Tech Solutions Ltd (Employer) and James O'Brien (Employee).

1. Position
Employee is appointed as Senior Software Engineer effective 1 May 2026.

2. Salary
Monthly gross salary of GBP 6,500, payable on the last business day of each month.

3. Working Hours
37.5 hours per week, Monday to Friday.

4. Annual Leave
25 days paid annual leave per year, plus public holidays.

5. Notice Period
During probation (6 months): 1 week notice by either party.
After probation: 3 months notice by either party.

6. Confidentiality
Employee shall keep all company information strictly confidential during and after employment.

7. Governing Law
This contract is governed by the laws of England and Wales.

Signed: 1 April 2026.
""",
)

make_txt(
    FIXTURES / "txt" / "12_high_risk_unilateral_id.txt",
    """PERJANJIAN LAYANAN

Antara PT Solusi Digital (Penyedia) dan CV Mitra Usaha (Klien).

1. Layanan
Penyedia akan menyediakan layanan pengembangan perangkat lunak.

2. Pembayaran
Klien wajib membayar Rp 50.000.000 per bulan.
Denda 15% per hari akan dikenakan atas keterlambatan pembayaran.
Bunga keterlambatan sebesar 20% per bulan akan dikenakan secara otomatis.

3. Perubahan Sepihak
Penyedia layanan dapat mengubah perjanjian ini kapan saja tanpa pemberitahuan.
Pihak pertama melepaskan semua hak hukum yang dimilikinya berdasarkan perjanjian ini.

4. Tanggung Jawab
Penyedia tidak bertanggung jawab atas kerugian apapun yang dialami Klien.

5. Penyelesaian Sengketa
[Klausul penyelesaian sengketa akan ditambahkan kemudian.]

Ditandatangani:
""",
)

make_txt(
    FIXTURES / "txt" / "13_medium_risk_lease_nl.txt",
    """HUUROVEREENKOMST

Tussen Van der Berg Vastgoed BV (Verhuurder) en Innovate NL BV (Huurder).

Artikel 1 - Partijen
Verhuurder: Van der Berg Vastgoed BV, Amsterdam. Huurder: Innovate NL BV, Utrecht.

Artikel 2 - Gehuurde
Kantoorruimte aan de Keizersgracht 200, Amsterdam, Nederland (BW).

Artikel 3 - Huurprijs
EUR 3.500 per maand, te betalen voor de 1e van elke maand.
Rente van 15% per maand wordt in rekening gebracht op achterstallige betalingen.

Artikel 4 - Duur
3 jaar vanaf 1 mei 2026.

Artikel 5 - Toepasselijk recht
Nederlands recht (Burgerlijk Wetboek).

Ondertekend: 1 april 2026.
""",
)

make_txt(
    FIXTURES / "txt" / "14_low_risk_nda_en.txt",
    """NON-DISCLOSURE AGREEMENT

Between Quantum Analytics Ltd (Disclosing Party) and DataSystems Inc (Receiving Party).

1. Purpose
The parties wish to explore a potential commercial partnership and may share confidential information.

2. Confidential Information
All business, technical, and financial data shared between the parties under this agreement.

3. Obligations
Receiving Party shall not disclose or use confidential information for any purpose other than
evaluating the potential partnership, without prior written consent.

4. Exceptions
Obligations do not apply to information that is publicly available or required by law to be disclosed.

5. Duration
This agreement is effective for 3 years from the date of signing.

6. Governing Law
This agreement is governed by the laws of England and Wales.
Any disputes shall be subject to the exclusive jurisdiction of the courts of England and Wales.

Signed: 1 April 2026.
""",
)

make_txt(
    FIXTURES / "txt" / "15_critical_risk_no_law_en.txt",
    """CONTRACT FOR SERVICES

Between Alpha Provider and Beta Client.

1. Services
Alpha Provider will deliver [services TBD].

2. Payment
Client must pay USD 5,000 per week.
A penalty of 50% per day applies for late payment.
Failure to pay within 3 days triggers a penalty of 50% of the total amount due.

3. Termination
Alpha Provider may terminate immediately and without cause at its sole discretion.
Client waives any right to compensation in the event of early termination by the provider.

4. Disputes
All disputes shall be resolved exclusively in favour of the service provider.

5. Liability
No liability whatsoever shall be incurred by the service provider under any circumstances.

6. Governing Law
[Governing law: TBD — to be confirmed by both parties at a later stage.]
[Jurisdiction: to be discussed and confirmed at a later date.]
""",
)

make_txt(
    FIXTURES / "txt" / "16_notice_id.txt",
    """PENGUMUMAN RESMI

Dengan ini diumumkan bahwa PT Sumber Makmur akan mengadakan Rapat Umum Pemegang Saham
Tahunan pada tanggal 15 Mei 2026 di Jakarta.

Agenda rapat meliputi:
- Laporan tahunan direksi
- Persetujuan laporan keuangan tahun 2025
- Pembagian dividen
- Penunjukan auditor independen

Para pemegang saham yang berhalangan hadir dapat memberikan kuasa kepada pihak lain
sesuai dengan ketentuan yang berlaku.

Jakarta, 1 April 2026
Direksi PT Sumber Makmur
""",
)


# ── Negative fixtures ─────────────────────────────────────────────────────────

# Empty file
(FIXTURES / "negative" / "empty.pdf").write_bytes(b"")

# Fake PDF (plain text with .pdf extension)
(FIXTURES / "negative" / "fake.pdf").write_bytes(
    b"This is not a real PDF. Just a text file renamed to .pdf."
)

# CSV file
(FIXTURES / "negative" / "test.csv").write_bytes(
    b"id,name,value\n1,Alpha,100\n2,Beta,200\n3,Gamma,300\n"
)

# Minimal valid PNG (1x1 red pixel)
_png_data = (
    b"\x89PNG\r\n\x1a\n"
    b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
    b"\x00\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd8N"
    b"\x00\x00\x00\x00IEND\xaeB`\x82"
)
(FIXTURES / "negative" / "test.png").write_bytes(_png_data)

print("Fixtures created:")
for p in sorted(FIXTURES.rglob("*")):
    if p.is_file():
        print(f"  {p.relative_to(FIXTURES)}  ({p.stat().st_size} bytes)")

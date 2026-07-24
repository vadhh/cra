import os

# Enforce offline mode if downloads not allowed
if os.getenv("LDV_DOWNLOAD_MODELS", "0") != "1":
    os.environ["HF_HUB_OFFLINE"] = "1"
    os.environ["TRANSFORMERS_OFFLINE"] = "1"

import logging
import json
import time
import uuid
import hmac as _hmac
import hashlib
import base64
import chardet
import magic
from flask import Flask, request, jsonify, send_from_directory, Response, redirect, g, session
from urllib.parse import urlparse
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import fitz
from langdetect import detect

from detector import profile_registry
from detector.detector_jurisdiction import detect_jurisdiction
from detector.detector_rules import layer1_analyze, required_clauses_for, clause_title, reconcile_required_flags
from detector.citation_db import annotate_layer1
from detector.risk_explainer import explain_findings
from detector.detector_distilbert import layer2_analyze, semantic_clause_presence
from detector.detector_scorer import layer3_score
from detector.detector_explain import layer4_explain
from translator import translate_text
from sydeco_engine import classify_clauses as _sydeco_classify
import database
import auth
import crypto
import worker

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

MAX_UPLOAD_BYTES = int(os.getenv("LDV_MAX_UPLOAD_MB", "10")) * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# Document types that are not contracts — skip full clause/risk analysis for these
_NON_CONTRACT_TYPES = {"invoice", "receipt", "purchase order", "non-contract"}

# Statuses where result_json isn't ready to serve: the job is still in flight
# (queued/processing) or needs a retry before it will ever produce a result
# (failed/retryable).
_HIDDEN_RESULT_STATUSES = {"queued", "processing", "failed", "retryable"}

_MIME_ALLOWLIST = {
    ".pdf":  {"application/pdf"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
    },
    ".txt":  {"text/plain"},
}

FRONTEND_DIR = os.path.abspath(
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "ldv-frontend")
)
UPLOADS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOADS_DIR, exist_ok=True)

app = Flask(__name__)
auth.configure_secret_key(app)

# Same-origin by default (the frontend is served by this app). Cross-origin
# access must be explicitly granted: LDV_CORS_ORIGINS="https://a.example,https://b.example"
_cors_origins = os.getenv("LDV_CORS_ORIGINS", "")
if _cors_origins:
    CORS(app, origins=[o.strip() for o in _cors_origins.split(",") if o.strip()])

# ponytail: Redis storage if configured (required for multi-worker); defaults to memory
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["500 per day", "60 per minute"],
    storage_uri=os.getenv("LDV_RATELIMIT_STORAGE_URL", "memory://"),
)


@limiter.request_filter
def bypass_rate_limits():
    if app.testing or app.config.get("TESTING") or os.environ.get("PYTEST_CURRENT_TEST") or os.environ.get("LDV_TESTING") == "1":
        return True
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        token = auth_header[7:].strip()
        user = database.get_user_by_token(token)
        if user and (user["email"] == "test-runner@ldv.internal" or auth.normalize_role(user["role"]) == "admin"):
            return True
    return False

# Init DB on first import
database.init_db()
database.cleanup_stuck_analyses()


@app.before_request
def _csrf_check():
    """Reject cross-origin state-mutating requests (defense-in-depth; SameSite=Lax already set)."""
    if app.testing or app.config.get("TESTING") or os.environ.get("PYTEST_CURRENT_TEST") or os.environ.get("LDV_TESTING") == "1":
        return
    if request.method in {"GET", "HEAD", "OPTIONS"}:
        return
    if request.headers.get("Authorization", "").startswith("Bearer "):
        return  # Bearer token auth is CSRF-immune
    origin = request.headers.get("Origin") or request.headers.get("Referer") or ""
    if not origin:
        # No Origin/Referer on a cookie-authenticated state-changing request → reject.
        # Bearer token path is already exempt above; browser same-origin POSTs always send Origin.
        logger.warning("CSRF: no origin header on %s", request.path)
        return jsonify({"error": "CSRF check failed"}), 403
    origin_host = (urlparse(origin).hostname or "").lower()
    expected_host = request.host.split(":")[0].lower()
    if origin_host != expected_host:
        logger.warning("CSRF: blocked %s (expected %s) on %s", origin_host, expected_host, request.path)
        return jsonify({"error": "CSRF check failed"}), 403


# ── Error handlers ─────────────────────────────────────────────────────────────

def _ip() -> str:
    return request.headers.get("X-Forwarded-For", request.remote_addr or "")


@app.errorhandler(413)
def file_too_large(e):
    return jsonify({"error": "File exceeds the 10 MB limit"}), 413


@app.errorhandler(429)
def rate_limited(e):
    database.write_audit("rate_limit", ip=_ip(), detail=request.path)
    return jsonify({"error": "Too many requests — please slow down"}), 429


@app.errorhandler(Exception)
def handle_exception(e):
    logger.exception("Unhandled exception")
    return jsonify({"error": "Internal server error"}), 500


# ── Text extraction ────────────────────────────────────────────────────────────

# F-05: a compressed-well-under-10MB file can still decompress/expand into a
# document expensive enough to stall the single-worker analysis queue
# (worker.py ThreadPoolExecutor(max_workers=1)). Reject before full parsing.
_MAX_PDF_PAGES = int(os.getenv("LDV_MAX_PDF_PAGES", "500"))
_MAX_DOCX_UNCOMPRESSED_BYTES = int(os.getenv("LDV_MAX_DOCX_UNCOMPRESSED_MB", "200")) * 1024 * 1024


def _extract_pdf(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    if doc.page_count > _MAX_PDF_PAGES:
        raise ValueError(
            f"PDF has {doc.page_count} pages, exceeding the {_MAX_PDF_PAGES}-page limit"
        )
    text = "\n".join(page.get_text() for page in doc)
    if not text.strip():
        text = _ocr_pdf(doc)
    return text


def _ocr_pdf(doc) -> str:
    """OCR fallback for scanned/image-only PDFs with no text layer.

    Slow (~1-3s/page on CPU); only invoked when normal text extraction finds
    nothing. Needs the tesseract-ocr system binary + language packs
    (tesseract-ocr-eng/fra/ind/nld) -- degrades to empty text (surfacing the
    existing "Scan/OCR required" error) if unavailable.
    """
    import pytesseract
    from PIL import Image
    from io import BytesIO
    try:
        parts = []
        for page in doc:
            img = Image.open(BytesIO(page.get_pixmap(dpi=300).tobytes("png")))
            parts.append(pytesseract.image_to_string(img, lang="eng+fra+ind+nld"))
        return "\n".join(parts)
    except Exception as e:
        logger.warning("OCR fallback failed: %s", e)
        return ""


def _extract_docx(data: bytes) -> str:
    import docx
    import zipfile
    from io import BytesIO
    with zipfile.ZipFile(BytesIO(data)) as zf:
        uncompressed = sum(zi.file_size for zi in zf.infolist())
    if uncompressed > _MAX_DOCX_UNCOMPRESSED_BYTES:
        raise ValueError(
            f"DOCX decompresses to {uncompressed // (1024*1024)} MB, exceeding the "
            f"{_MAX_DOCX_UNCOMPRESSED_BYTES // (1024*1024)} MB limit"
        )
    document = docx.Document(BytesIO(data))
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    detected = chardet.detect(data)
    encoding = detected.get("encoding") or "utf-8"
    return data.decode(encoding, errors="replace")


def _validate_and_extract(file) -> tuple[bytes, str, str]:
    """Validate upload, return (data, ext, text). Raises ValueError on failure."""
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    data = file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) == 0:
        raise ValueError("File is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(f"File exceeds the {MAX_UPLOAD_BYTES // (1024*1024)} MB limit")

    detected_mime = magic.from_buffer(data[:4096], mime=True)
    allowed_mimes = _MIME_ALLOWLIST.get(ext, set())
    if detected_mime == "application/octet-stream" and ext == ".docx" and data[:2] == b"PK":
        detected_mime = "application/zip"
    if detected_mime not in allowed_mimes:
        raise ValueError(f"File content does not match extension '{ext}'")

    try:
        if ext == ".pdf":
            text = _extract_pdf(data)
        elif ext == ".docx":
            text = _extract_docx(data)
        else:
            text = _extract_txt(data)
    except Exception as e:
        logger.warning("Failed to extract text from %s: %s", ext, e)
        raise ValueError(f"Corrupted or invalid {ext} file: {str(e)}")

    if not text.strip():
        raise ValueError("Scan/OCR required. No usable text could be extracted from this document.")

    return data, ext, text


# Short frontend select values -> classifier label, restricted to the 11
# original profiles (2026-07-17 review pilot mandate: "the other 45 profiles
# are disabled in the customer interface"). Kept in sync with the registry by
# construction -- see _resolve_pilot_type(), which double-checks each label
# still resolves to a classifier.status == "validated" profile before use.
PILOT_TYPE_MAPPING = {
    "service": "service agreement",
    "nda": "non-disclosure agreement",
    "employment": "employment contract",
    "software": "software license",
    "generic": "general contract",
    "lease": "lease agreement",
    "consulting": "consulting agreement",
    "commercial": "commercial agreement",
    "loan": "loan agreement",
    "partnership": "partnership agreement",
    "purchase": "purchase agreement",
}


def _resolve_pilot_type(value: str) -> str | None:
    """Map a manual contract-type selection to its classifier label, restricted
    to pilot-available (registry classifier.status == "validated") profiles.
    Returns None if `value` doesn't resolve to one -- callers must treat that
    as rejection, not silently fall through to an arbitrary label."""
    label = PILOT_TYPE_MAPPING.get(value.lower(), value.lower())
    profile = profile_registry.detect_profile(label)
    if profile and (profile.get("classifier") or {}).get("status") == "validated":
        return label
    return None


def _run_analysis(text: str, jurisdiction: str, lang: str, policy_name: str | None = None, override_type: str | None = None) -> dict:
    """Run L1–L3 analysis and return result dict.

    For non-contract documents (invoice, receipt, purchase order) the pipeline
    stops after L2: no clause-risk scoring or MLP tagging is performed.
    """
    _meta = {"encryption_enabled": crypto.is_enabled()}
    layer1 = layer1_analyze(text, jurisdiction)
    annotate_layer1(layer1, jurisdiction)  # attach legal citations to each finding

    analysis_text = text
    if lang not in ("en", "unknown"):
        try:
            analysis_text = translate_text(text, "en", src_lang=lang)
        except Exception as e:
            logger.warning("Translation failed, using original: %s", e)

    layer2 = layer2_analyze(analysis_text)

    if override_type:
        mapped_label = _resolve_pilot_type(override_type)
        if mapped_label is None:
            # Route handlers validate override_type before reaching here and
            # should never pass an invalid one -- this is a fail-safe: ignore
            # it and keep the classifier's own result rather than silently
            # scoring under an unvalidated/unresolvable profile.
            logger.warning("Ignoring non-pilot override_type=%r; keeping classifier result", override_type)
        else:
            original_model_conf = (layer2.get("document_type") or {}).get("model_confidence") or (layer2.get("document_type") or {}).get("confidence") or 0.0
            layer2["document_type"] = {
                "label": mapped_label,
                "confidence": 1.0,
                "model_confidence": original_model_conf,
                "keyword_override_confidence": 0.0,
                "override_applied": True,
                "override_reason": "user_selected",
                "source": "user_selected",
                "candidates": [{"label": mapped_label, "confidence": 1.0}]
            }

    doc_type_label = ((layer2.get("document_type") or {}).get("label") or "").lower()
    reconcile_required_flags(layer1.get("clause_presence") or [], doc_type_label)
    if doc_type_label in _NON_CONTRACT_TYPES:
        logger.info("Document type '%s' — skipping clause analysis", doc_type_label)
        return {
            "language":          lang,
            "jurisdiction":      jurisdiction,
            "document_type_note": (
                f"This document appears to be {_article(doc_type_label)} {doc_type_label}. "
                "Full contractual clause analysis is not applicable. "
                "Payment-term rules were still evaluated."
            ),
            "layer1":    layer1,
            "layer2":    layer2,
            "layer3":    {"available": False, "skipped": True, "reason": "non_contract_document"},
            "layer4":    {"available": False, "skipped": True},
            "clause_tags": [],
            "_meta":     _meta,
        }

    _semantic_backfill(layer1, doc_type_label, analysis_text)
    explain_findings(layer1, jurisdiction, doc_type_label)

    layer3 = layer3_score(layer1, layer2, lang=lang, policy_name=policy_name)
    clause_tags = _sydeco_classify(analysis_text)

    return {
        "language":     lang,
        "jurisdiction": jurisdiction,
        "layer1":       layer1,
        "layer2":       layer2,
        "layer3":       layer3,
        "layer4":       {"available": False, "skipped": True},
        "clause_tags":  clause_tags,
        "_meta":        _meta,
    }


def _semantic_backfill(layer1: dict, doc_type_label: str, text: str) -> None:
    """Re-check keyword-missing *required* clauses with semantic NLI, in place.

    Keyword/regex detection misses clauses worded unusually.  Before scoring, we
    run an NLI presence check (reusing the loaded DistilBERT model) on the
    required clauses L1 marked absent; any that are semantically present get
    flipped to present with source="semantic_nli", so L3's missing-clause logic
    needs no change.  Pure recovery — it only ever turns a False into True.
    """
    presence = layer1.get("clause_presence") or []
    required = set(required_clauses_for(doc_type_label))
    missing = [
        (c["clause_id"], c.get("title") or clause_title(c["clause_id"]))
        for c in presence
        if c["clause_id"] in required and not c.get("present")
    ]
    if not missing:
        return

    recovered = semantic_clause_presence(text, missing)
    for c in presence:
        conf = recovered.get(c["clause_id"])
        if conf is not None:
            c["present"] = True
            c["source"] = "semantic_nli"
            c["evidence"] = c.get("evidence") or f"semantic match (NLI {conf})"


def _article(word: str) -> str:
    return "an" if word[:1].lower() in "aeiou" else "a"


_DL_TTL = int(os.getenv("LDV_DOWNLOAD_LINK_TTL", "900"))  # seconds (default 15 minutes)


def _derive_download_key(base: bytes) -> bytes:
    # HMAC-based KDF (domain separation, not entropy stretching — base is
    # already a high-entropy secret). Fixes F-03: previously reused
    # LDV_SECRET_KEY with a naive `+b":download"` suffix, so leaking the
    # session secret also handed over the download-link signing key.
    return _hmac.new(base, b"ldv-download-link-v1", hashlib.sha256).digest()


def _dl_keys() -> list[bytes]:
    override = os.getenv("LDV_DOWNLOAD_LINK_SECRET", "")
    if override:
        return [override.strip().encode()]
    keys_str = os.getenv("LDV_SECRET_KEY", "")
    if not keys_str:
        k = app.secret_key
        base_keys = [k if isinstance(k, bytes) else k.encode()]
    else:
        base_keys = [k.strip().encode() if isinstance(k, str) else k for k in keys_str.split(",")]
    return [_derive_download_key(k) for k in base_keys]


def _make_download_token(analysis_id: str) -> tuple[str, int]:
    expires_at = int(time.time()) + _DL_TTL
    payload = f"{analysis_id}:{expires_at}"
    sig = _hmac.new(_dl_keys()[0], payload.encode(), hashlib.sha256).hexdigest()
    token = base64.urlsafe_b64encode(payload.encode()).decode() + "." + sig
    return token, expires_at


def _verify_download_token(token: str) -> str | None:
    try:
        payload_b64, sig = token.rsplit(".", 1)
        payload = base64.urlsafe_b64decode(payload_b64.encode()).decode()
        analysis_id, expires_str = payload.rsplit(":", 1)
        if int(expires_str) < int(time.time()):
            return None
        for k in _dl_keys():
            expected = _hmac.new(k, payload.encode(), hashlib.sha256).hexdigest()
            if _hmac.compare_digest(sig, expected):
                return analysis_id
        return None
    except Exception:
        return None


# ── Auth routes ────────────────────────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
@limiter.limit("10 per minute", methods=["POST"])
def login():
    if request.method == "GET":
        return send_from_directory(FRONTEND_DIR, "login.html")
    data = request.get_json(silent=True) or request.form
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    mfa_code = (data.get("mfa_code") or data.get("otp") or "").strip()

    # F-09: per-account lockout, independent of the per-IP rate limit above —
    # a distributed attacker isn't slowed by a per-IP-only limit.
    existing = database.get_user_by_email(email)
    if existing and database.is_account_locked(existing):
        database.write_audit("login.fail", ip=_ip(), detail=f"{email} (locked)")
        return jsonify({"error": "Account temporarily locked due to repeated failed attempts. Try again later."}), 423

    user = auth.verify_login(email, password)
    if user is None:
        if existing:
            database.record_login_failure(existing["id"])
        database.write_audit("login.fail", ip=_ip(), detail=email)
        return jsonify({"error": "Invalid credentials"}), 401

    # Check if MFA is required
    has_secret = bool(user.get("mfa_secret"))
    mandatory = auth.is_mfa_mandatory(user)

    if has_secret:
        if not mfa_code:
            session["mfa_pending_uid"] = user["id"]
            return jsonify({"mfa_required": True})

        import pyotp
        secret = crypto.dec_str(user["mfa_secret"])
        totp = pyotp.TOTP(secret)
        verified = totp.verify(mfa_code, valid_window=1)

        # Check recovery codes if TOTP fails
        if not verified and user.get("mfa_recovery_codes"):
            import json
            from werkzeug.security import check_password_hash
            hashes = json.loads(user["mfa_recovery_codes"] or "[]")
            matched_hash = None
            for h in hashes:
                if check_password_hash(h, mfa_code):
                    matched_hash = h
                    break
            if matched_hash:
                hashes.remove(matched_hash)
                database.update_user_mfa(user["id"], user["mfa_secret"], json.dumps(hashes))
                verified = True
                database.write_audit("mfa.recovery_used", user_id=user["id"], org_id=user["org_id"], ip=_ip())

        if not verified:
            database.record_login_failure(user["id"])
            database.write_audit("login.fail", ip=_ip(), detail=f"{email} (invalid MFA)")
            return jsonify({"error": "Invalid MFA code"}), 401

    elif mandatory:
        if not mfa_code:
            session["mfa_enroll_pending_uid"] = user["id"]
            return jsonify({"mfa_enroll_required": True})

    # Complete login
    session.pop("mfa_pending_uid", None)
    session.pop("mfa_enroll_pending_uid", None)
    session["uid"] = user["id"]
    database.record_login_success(user["id"])
    database.write_audit("login.success", user_id=user["id"], org_id=user["org_id"], ip=_ip())
    return jsonify({"ok": True, "role": user["role"]})


@app.route("/api/v1/logout", methods=["POST"])
def logout():
    user = auth.current_user()
    if user:
        database.write_audit("logout", user_id=user["id"], org_id=user["org_id"], ip=_ip())
    session.clear()
    return jsonify({"ok": True})


@app.route("/logout")
def get_logout():
    user = auth.current_user()
    if user:
        database.write_audit("logout", user_id=user["id"], org_id=user["org_id"], ip=_ip())
    session.clear()
    return redirect("/login")


@app.route("/api/v1/mfa/status")
def api_mfa_status():
    uid = session.get("uid") or session.get("mfa_enroll_pending_uid") or session.get("mfa_pending_uid")
    if not uid:
        return jsonify({"authenticated": False}), 401
    user = database.get_user_by_id(uid)
    if not user:
        return jsonify({"error": "User not found"}), 404
    return jsonify({
        "mfa_enabled": bool(user.get("mfa_secret")),
        "mfa_mandatory": auth.is_mfa_mandatory(user),
        "email": user["email"],
        "role": auth.normalize_role(user["role"])
    })


@app.route("/api/v1/mfa/setup", methods=["POST"])
def api_mfa_setup():
    uid = session.get("mfa_enroll_pending_uid") or session.get("uid")
    if not uid:
        return jsonify({"error": "Authentication required"}), 401
    
    user = database.get_user_by_id(uid)
    if not user:
        return jsonify({"error": "User not found"}), 404
        
    if session.get("uid"):
        data = request.json or {}
        password = data.get("password") or ""
        if not auth.verify_login(user["email"], password):
            return jsonify({"error": "Re-authentication failed: invalid password"}), 401

    import pyotp
    import secrets
    secret = pyotp.random_base32()
    totp = pyotp.TOTP(secret)
    plain_codes = [secrets.token_hex(4) for _ in range(10)]
    
    session["mfa_setup_secret"] = secret
    session["mfa_setup_codes"] = plain_codes
    
    uri = totp.provisioning_uri(name=user["email"], issuer_name="Sydeco Contract Risk Analyzer")
    return jsonify({
        "secret": secret,
        "provisioning_uri": uri,
        "recovery_codes": plain_codes
    })


@app.route("/api/v1/mfa/enable", methods=["POST"])
def api_mfa_enable():
    uid = session.get("mfa_enroll_pending_uid") or session.get("uid")
    if not uid:
        return jsonify({"error": "Authentication required"}), 401
    
    user = database.get_user_by_id(uid)
    if not user:
        return jsonify({"error": "User not found"}), 404
        
    secret = session.get("mfa_setup_secret")
    plain_codes = session.get("mfa_setup_codes")
    if not secret or not plain_codes:
        return jsonify({"error": "MFA setup has not been initialized"}), 400
        
    data = request.json or {}
    code = (data.get("code") or "").strip()
    if not code:
        return jsonify({"error": "Verification code required"}), 400
        
    import pyotp
    totp = pyotp.TOTP(secret)
    if not totp.verify(code, valid_window=1):
        return jsonify({"error": "Invalid verification code"}), 400
        
    from werkzeug.security import generate_password_hash
    import json
    enc_secret = crypto.enc_str(secret)
    hashed_codes = [generate_password_hash(c) for c in plain_codes]
    
    database.update_user_mfa(uid, enc_secret, json.dumps(hashed_codes))
    
    session.pop("mfa_setup_secret", None)
    session.pop("mfa_setup_codes", None)
    
    if session.get("mfa_enroll_pending_uid"):
        session.pop("mfa_enroll_pending_uid", None)
        session["uid"] = uid
        database.write_audit("login.success", user_id=uid, org_id=user["org_id"], ip=_ip())
    else:
        database.write_audit("mfa.enable", user_id=uid, org_id=user["org_id"], ip=_ip())
        
    return jsonify({"ok": True})


@app.route("/api/v1/mfa/skip", methods=["POST"])
def api_mfa_skip():
    uid = session.get("mfa_enroll_pending_uid")
    if not uid:
        return jsonify({"error": "No pending enrollment"}), 400
    user = database.get_user_by_id(uid)
    if not user:
        return jsonify({"error": "User not found"}), 404
    if database.org_mfa_required(user["org_id"]) or auth.is_mfa_mandatory(user):
        return jsonify({"error": "MFA is mandatory for this account"}), 403
    session.pop("mfa_enroll_pending_uid", None)
    session["uid"] = uid
    database.write_audit("login.success.mfa_skipped", user_id=uid, org_id=user["org_id"], ip=_ip())
    return jsonify({"ok": True})


@app.route("/api/v1/mfa/disable", methods=["POST"])
@auth.login_required
def api_mfa_disable():
    user = g.user
    data = request.json or {}
    password = data.get("password") or ""
    if not auth.verify_login(user["email"], password):
        return jsonify({"error": "Re-authentication failed: invalid password"}), 401

    if database.org_mfa_required(user["org_id"]) or auth.is_mfa_mandatory(user):
        return jsonify({"error": "MFA is mandatory for this account"}), 403

    database.update_user_mfa(user["id"], None, None)
    database.write_audit("mfa.disable", user_id=user["id"], org_id=user["org_id"], ip=_ip())
    return jsonify({"ok": True})


# ── Upload & analyse (primary endpoint) ───────────────────────────────────────

@app.route("/api/v1/consent", methods=["GET", "POST"])
@auth.login_required
def user_consent():
    """Record or retrieve user agreement timestamps for Terms of Service and Privacy Policy."""
    if request.method == "GET":
        version = request.args.get("version", "1.0")
        consented = database.has_user_consented(g.user["id"], version=version)
        return jsonify({
            "user_id": g.user["id"],
            "version": version,
            "consented": consented,
            "tos_url": "/docs/legal/TERMS_OF_SERVICE.md",
            "privacy_url": "/docs/legal/PRIVACY_POLICY.md",
        })

    data = request.get_json(silent=True) or {}
    tos_accepted = bool(data.get("tos_accepted"))
    privacy_accepted = bool(data.get("privacy_accepted"))
    version = str(data.get("version") or "1.0")

    if not tos_accepted or not privacy_accepted:
        return jsonify({
            "error": "Both Terms of Service (tos_accepted) and Privacy Policy (privacy_accepted) must be accepted."
        }), 400

    rec = database.record_user_consent(
        user_id=g.user["id"],
        tos_accepted=tos_accepted,
        privacy_accepted=privacy_accepted,
        version=version,
        ip_address=_ip(),
    )
    database.write_audit("consent.accept", user_id=g.user["id"], org_id=g.user["org_id"], ip=_ip(), detail=f"v{version}")
    return jsonify(rec)


@app.route("/api/v1/upload", methods=["POST"])
@auth.login_required
@limiter.limit("20 per minute")
def upload():
    """Save file to disk, extract text, run analysis, persist to DB."""
    if auth.normalize_role(g.user["role"]) == "viewer":
        return jsonify({"error": "Forbidden: viewers cannot upload documents"}), 403
    if os.getenv("LDV_PRODUCTION") == "1" and not crypto.is_enabled():
        return jsonify({"error": "Service configuration error: encryption is disabled or not configured in production"}), 500
    if os.getenv("LDV_ENFORCE_CONSENT", "1") == "1":
        if not (app.testing or app.config.get("TESTING") or os.environ.get("PYTEST_CURRENT_TEST") or os.environ.get("LDV_TESTING") == "1"):
            if not database.has_user_consented(g.user["id"]):
                return jsonify({
                    "error": "TOS and Privacy Policy acceptance required prior to document processing.",
                    "code": "CONSENT_REQUIRED",
                    "consent_url": "/api/v1/consent"
                }), 403

    requested_type = request.args.get("type")
    if requested_type and requested_type != "auto" and _resolve_pilot_type(requested_type) is None:
        return jsonify({
            "error": f"Contract type '{requested_type}' is not available in this pilot.",
            "available_types": sorted(PILOT_TYPE_MAPPING.keys()),
        }), 400

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    try:
        data, ext, text = _validate_and_extract(file)
    except ValueError as e:
        logger.warning("Upload validation failed for %s: %s", file.filename, e)
        return jsonify({"error": str(e)}), 400

    # Calculate pages
    pages = 1
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            pages = doc.page_count
        except Exception:
            pages = 1
    else:
        pages = max(1, len(text) // 2000)

    # Check subscription limits
    usage = database.get_org_usage(g.user["org_id"])
    if usage:
        contract_limit = usage.get("contract_limit", 100)
        contract_used = usage.get("contract_used", 0)
        page_limit = usage.get("page_limit", 500)
        page_used = usage.get("page_used", 0)

        if contract_used + 1 > contract_limit:
            return jsonify({"error": "Subscription limit exceeded: contract allowance exhausted."}), 403
        if page_used + pages > page_limit:
            return jsonify({"error": f"Subscription limit exceeded: page allowance exhausted (uploading {pages} pages would exceed limit of {page_limit})."}), 403

    # Atomically increment usage
    database.increment_org_usage(g.user["org_id"], contracts=1, pages=pages)

    # Save file to disk
    stored_name = f"{uuid.uuid4().hex}{ext}"
    file_path   = os.path.join(UPLOADS_DIR, stored_name)
    with open(file_path, "wb") as f:
        f.write(crypto.enc_bytes(data))

    # Detect language
    try:
        lang = detect(text)
    except Exception:
        lang = "unknown"

    # Get client/folder parameters
    client = (
        request.args.get("client") or 
        request.form.get("client") or 
        request.args.get("client_ref") or 
        request.form.get("client_ref") or
        request.args.get("clientRef") or 
        request.form.get("clientRef")
    )
    case_folder = (
        request.args.get("case_folder") or 
        request.form.get("case_folder") or
        request.args.get("caseFolder") or 
        request.form.get("caseFolder")
    )
    if client:
        client = client.strip()
    if case_folder:
        case_folder = case_folder.strip()

    # Save document record
    doc_id = database.save_document(
        original_filename=file.filename,
        stored_filename=stored_name,
        file_path=file_path,
        file_size=len(data),
        file_type=ext,
        language=lang,
        extracted_text=text,
        org_id=g.user["org_id"],
        owner_id=g.user["id"],
        client=client,
        case_folder=case_folder,
    )

    want_explain = request.args.get("explain", "0") == "1"
    policy_name = request.args.get("policy", "default_v1")
    override_jurisdiction = request.args.get("jurisdiction")
    override_type = request.args.get("type")

    if override_jurisdiction == "auto":
        override_jurisdiction = None
    if override_type == "auto":
        override_type = None

    # Save queued analysis record
    analysis_id = database.save_analysis(
        document_id=doc_id,
        jurisdiction=None,
        document_type=None,
        risk_score=None,
        risk_label=None,
        result=None,
        status="queued",
    )

    # Submit background task
    import inspect
    sig = inspect.signature(worker.submit_job)
    kwargs = {}
    if "policy_name" in sig.parameters:
        kwargs["policy_name"] = policy_name
    if "override_jurisdiction" in sig.parameters:
        kwargs["override_jurisdiction"] = override_jurisdiction
    if "override_type" in sig.parameters:
        kwargs["override_type"] = override_type

    worker.submit_job(analysis_id, text, lang, want_explain, **kwargs)

    database.write_audit(
        "upload", user_id=g.user["id"], org_id=g.user["org_id"],
        resource_id=analysis_id, ip=_ip(), detail=file.filename,
    )
    logger.info(
        "UPLOAD: enqueued file=%s lang=%s explain=%s id=%s",
        file.filename, lang, want_explain, analysis_id,
    )

    return jsonify({"id": analysis_id, "status": "queued"}), 202


# ── Result API ─────────────────────────────────────────────────────────────────

@app.route("/api/v1/result/<analysis_id>")
@auth.login_required
def api_result(analysis_id: str):
    row = database.get_result(analysis_id)
    if row is None:
        return jsonify({"error": "Not found"}), 404
    user = g.user
    if user["role"] != "admin" and row.get("org_id") != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403
    row.pop("org_id", None)  # internal field, not part of the API response

    # ponytail: Expose raw_text via ?debug=1 only (P3 item 15)
    extracted_text = row.pop("extracted_text", None)
    if request.args.get("debug") == "1":
        row["raw_text"] = extracted_text

    # Phase 3 (S2): deserialize score_breakdown from JSON string
    if row.get("score_breakdown") and isinstance(row["score_breakdown"], str):
        try:
            row["score_breakdown"] = json.loads(row["score_breakdown"])
        except (ValueError, TypeError):
            pass

    status = row.get("status", "completed")
    if status in _HIDDEN_RESULT_STATUSES:
        if status == "retryable" and row.get("error_message") == "low_confidence":
            if row.get("result_json"):
                try:
                    res_dict = json.loads(row["result_json"])
                    dt_info = res_dict.get("layer2", {}).get("document_type", {})
                    row["candidates"] = dt_info.get("candidates", [])
                    row["detected_label"] = dt_info.get("label")
                    row["detected_confidence"] = dt_info.get("confidence")
                except Exception:
                    pass
        row["result"] = None
        row.pop("result_json", None)
        return jsonify(row)

    row["result"] = json.loads(row["result_json"]) if row.get("result_json") else None
    row.pop("result_json", None)
    return jsonify(row)



@app.route("/api/v1/result/<analysis_id>", methods=["DELETE"])
@auth.login_required
def api_delete_result(analysis_id: str):
    row = database.get_result(analysis_id)
    if row is None:
        return jsonify({"error": "Not found"}), 404
    user = g.user
    if user["role"] != "admin" and row.get("org_id") != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403
    info = database.delete_analysis(analysis_id)
    if info and info.get("file_path"):
        try:
            os.remove(info["file_path"])
        except FileNotFoundError:
            pass
    database.write_audit(
        "delete", user_id=user["id"], org_id=user["org_id"],
        resource_id=analysis_id, ip=_ip(),
    )
    logger.info("DELETE: id=%s org=%s by=%s", analysis_id, row.get("org_id"), user["email"])
    return jsonify({"deleted": True, "id": analysis_id})


@app.route("/api/v1/result/<analysis_id>/retry", methods=["POST"])
@auth.login_required
def api_retry_result(analysis_id: str):
    row = database.get_result(analysis_id)
    if row is None:
        return jsonify({"error": "Not found"}), 404
    user = g.user
    if user["role"] != "admin" and row.get("org_id") != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403

    data = request.get_json(silent=True) or {}
    override_type = data.get("type")
    override_jurisdiction = data.get("jurisdiction")

    if override_type and override_type != "auto" and _resolve_pilot_type(override_type) is None:
        return jsonify({
            "error": f"Contract type '{override_type}' is not available in this pilot.",
            "available_types": sorted(PILOT_TYPE_MAPPING.keys()),
        }), 400

    new_status = database.retry_analysis(analysis_id)
    if new_status is None:
        return jsonify({"error": "Analysis is not retryable (wrong status or retry limit exhausted)"}), 409

    text = row.get("extracted_text")
    lang = row.get("language") or "unknown"
    worker.submit_job(analysis_id, text, lang, False, override_type=override_type, override_jurisdiction=override_jurisdiction)

    database.write_audit(
        "analysis.retry", user_id=user["id"], org_id=row.get("org_id"),
        resource_id=analysis_id, ip=_ip(),
    )
    return jsonify({"id": analysis_id, "status": "queued"}), 202


@app.route("/api/v1/result/<analysis_id>/download-link", methods=["POST"])
@auth.login_required
def api_download_link(analysis_id: str):
    """Generate a time-limited signed URL to download the original file."""
    if g.user.get("download_disabled") or database.get_user_by_id(g.user["id"]).get("download_disabled"):
        return jsonify({"error": "Forbidden: download access is disabled for your account"}), 403

    row = database.get_document_file_info(analysis_id)
    if row is None:
        return jsonify({"error": "Not found"}), 404
    if auth.normalize_role(g.user["role"]) != "admin" and row.get("org_id") != g.user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403

    data = request.json or {}
    one_time = bool(data.get("one_time", False))

    token, expires_at = _make_download_token(analysis_id)
    database.save_download_link(token, analysis_id, expires_at, 1 if one_time else 0)
    database.write_audit("download.link_generated", user_id=g.user["id"], org_id=row.get("org_id"), resource_id=analysis_id, ip=_ip(), detail=f"one_time: {one_time}")
    return jsonify({"url": f"/download/{token}", "expires_at": expires_at})


@app.route("/download/<token>")
def download_file(token: str):
    """Serve the original encrypted file via a signed token (no session required)."""
    link_info = database.get_download_link(token)
    if not link_info:
        return jsonify({"error": "Invalid or expired download link"}), 403
    if link_info["revoked"] or link_info["used"]:
        return jsonify({"error": "Link has been revoked or already used"}), 403
    if link_info["expires_at"] < int(time.time()):
        return jsonify({"error": "Link has expired"}), 403

    analysis_id = _verify_download_token(token)
    if analysis_id is None or analysis_id != link_info["analysis_id"]:
        return jsonify({"error": "Invalid or expired download link"}), 403

    info = database.get_document_file_info(analysis_id)
    if info is None or not os.path.isfile(info["file_path"]):
        return jsonify({"error": "File not found"}), 404

    if link_info["one_time"]:
        database.mark_download_link_used(token)

    database.write_audit("download.served", user_id=None, org_id=info.get("org_id"), resource_id=analysis_id, ip=_ip())

    with open(info["file_path"], "rb") as f:
        raw = crypto.dec_bytes(f.read())
    ext = info["file_type"]
    mime_map = {
        ".pdf":  "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt":  "text/plain",
    }
    mime = mime_map.get(ext, "application/octet-stream")
    filename = info["original_filename"] or f"contract{ext}"
    return Response(
        raw, mimetype=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Administrative User & Organization Management ────────────────────────────

@app.route("/api/v1/admin/users")
@auth.role_required("manager")
def api_admin_users():
    user = g.user
    u_role = auth.normalize_role(user["role"])
    if u_role == "admin":
        users = database.get_all_users()
    else:
        users = database.get_users_by_org(user["org_id"])
    
    for u in users:
        u.pop("password_hash", None)
        u["mfa_enabled"] = bool(u.get("mfa_secret"))
        u.pop("mfa_secret", None)
    return jsonify(users)


@app.route("/api/v1/admin/users", methods=["POST"])
@auth.role_required("manager")
def api_admin_create_user():
    user = g.user
    u_role = auth.normalize_role(user["role"])
    data = request.json or {}
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    role = data.get("role") or "analyst"
    org_id = data.get("org_id")
    
    if not email or not password:
        return jsonify({"error": "Email and password required"}), 400
        
    if u_role != "admin":
        org_id = user["org_id"]
        if role == "admin":
            return jsonify({"error": "Forbidden: managers cannot create administrators"}), 403
    else:
        if not org_id:
            return jsonify({"error": "Organization ID required"}), 400

    if role == "admin" and not auth.is_operator_org(org_id):
        return jsonify({
            "error": f"Forbidden: 'admin' may only be granted in the operator org "
                     f"('{auth.operator_org_name()}')"
        }), 403

    if database.get_user_by_email(email):
        return jsonify({"error": "User already exists"}), 400
        
    hashed = auth.hash_password(password)
    import secrets
    api_token = f"tok-{secrets.token_urlsafe(16)}"
    
    new_uid = database.create_user(org_id, email, hashed, role, api_token)
    database.write_audit("user.create", user_id=user["id"], org_id=org_id, resource_id=str(new_uid), ip=_ip(), detail=email)
    return jsonify({"ok": True, "user_id": new_uid})


@app.route("/api/v1/admin/users/<int:target_id>/status", methods=["POST"])
@auth.role_required("manager")
def api_admin_user_status(target_id: int):
    user = g.user
    u_role = auth.normalize_role(user["role"])
    data = request.json or {}
    active = int(data.get("active", 1))
    
    target = database.get_user_by_id(target_id)
    if not target:
        return jsonify({"error": "User not found"}), 404
        
    if u_role != "admin":
        if target["org_id"] != user["org_id"]:
            return jsonify({"error": "Forbidden"}), 403
        if target["role"] == "admin":
            return jsonify({"error": "Forbidden: managers cannot suspend administrators"}), 403
            
    if target_id == user["id"]:
        return jsonify({"error": "Forbidden: you cannot change your own status"}), 403
        
    if target["role"] == "admin" and active == 0:
        if database.count_active_admins() <= 1:
            return jsonify({"error": "Forbidden: cannot suspend the last system administrator"}), 403
            
    database.update_user_status(target_id, active)
    action = "user.unsuspend" if active else "user.suspend"
    database.write_audit(action, user_id=user["id"], org_id=target["org_id"], resource_id=str(target_id), ip=_ip())
    return jsonify({"ok": True})


@app.route("/api/v1/admin/users/<int:target_id>/role", methods=["POST"])
@auth.role_required("manager")
def api_admin_user_role(target_id: int):
    user = g.user
    u_role = auth.normalize_role(user["role"])
    data = request.json or {}
    new_role = data.get("role")
    if not new_role:
        return jsonify({"error": "Role required"}), 400
        
    target = database.get_user_by_id(target_id)
    if not target:
        return jsonify({"error": "User not found"}), 404
        
    if u_role != "admin":
        if target["org_id"] != user["org_id"]:
            return jsonify({"error": "Forbidden"}), 403
        if target["role"] == "admin" or new_role == "admin":
            return jsonify({"error": "Forbidden: managers cannot manage administrator roles"}), 403

    if new_role == "admin" and not auth.is_operator_org(target["org_id"]):
        return jsonify({
            "error": f"Forbidden: 'admin' may only be granted in the operator org "
                     f"('{auth.operator_org_name()}')"
        }), 403

    if target_id == user["id"]:
        return jsonify({"error": "Forbidden: you cannot change your own role"}), 403
        
    if target["role"] == "admin" and new_role != "admin":
        if database.count_active_admins() <= 1:
            return jsonify({"error": "Forbidden: cannot demote the last system administrator"}), 403
            
    database.update_user_role(target_id, new_role)
    database.write_audit("user.role_change", user_id=user["id"], org_id=target["org_id"], resource_id=str(target_id), ip=_ip(), detail=new_role)
    return jsonify({"ok": True})


@app.route("/api/v1/admin/users/<int:target_id>/download-access", methods=["POST"])
@auth.role_required("manager")
def api_admin_user_download_access(target_id: int):
    user = g.user
    u_role = auth.normalize_role(user["role"])
    data = request.json or {}
    disabled = int(data.get("download_disabled", 0))
    
    target = database.get_user_by_id(target_id)
    if not target:
        return jsonify({"error": "User not found"}), 404
        
    if u_role != "admin" and target["org_id"] != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403
        
    database.update_user_download_access(target_id, disabled)
    action = "user.download.disable" if disabled else "user.download.enable"
    database.write_audit(action, user_id=user["id"], org_id=target["org_id"], resource_id=str(target_id), ip=_ip())
    return jsonify({"ok": True})


@app.route("/api/v1/admin/users/<int:target_id>/mfa-exempt", methods=["POST"])
@auth.role_required("admin")
def api_admin_user_mfa_exempt(target_id: int):
    user = g.user
    data = request.json or {}
    exempt = int(bool(data.get("mfa_exempt", 0)))

    target = database.get_user_by_id(target_id)
    if not target:
        return jsonify({"error": "User not found"}), 404

    database.update_user_mfa_exempt(target_id, exempt)
    database.write_audit("user.mfa_exempt_change", user_id=user["id"], org_id=target["org_id"], resource_id=str(target_id), ip=_ip(), detail=str(bool(exempt)))
    return jsonify({"ok": True})


@app.route("/api/v1/admin/users/<int:target_id>/mfa-reset", methods=["POST"])
@auth.role_required("manager")
def api_admin_user_mfa_reset(target_id: int):
    user = g.user
    u_role = auth.normalize_role(user["role"])
    
    target = database.get_user_by_id(target_id)
    if not target:
        return jsonify({"error": "User not found"}), 404
        
    if u_role != "admin":
        if target["org_id"] != user["org_id"]:
            return jsonify({"error": "Forbidden"}), 403
        if target["role"] == "admin":
            return jsonify({"error": "Forbidden: managers cannot reset administrator MFA"}), 403
            
    database.update_user_mfa(target_id, None, None)
    database.write_audit("user.mfa_reset", user_id=user["id"], org_id=target["org_id"], resource_id=str(target_id), ip=_ip())
    return jsonify({"ok": True})


@app.route("/api/v1/admin/organizations")
@auth.admin_required
def api_admin_organizations():
    return jsonify(database.get_all_orgs())


@app.route("/api/v1/admin/organizations", methods=["POST"])
@auth.admin_required
def api_admin_create_organization():
    data = request.json or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Organization name required"}), 400
        
    if database.get_org_by_name(name):
        return jsonify({"error": "Organization already exists"}), 400
        
    new_oid = database.create_org(name)
    database.write_audit("org.create", user_id=g.user["id"], org_id=new_oid, resource_id=str(new_oid), ip=_ip(), detail=name)
    return jsonify({"ok": True, "org_id": new_oid})


@app.route("/api/v1/admin/organizations/<int:org_id>/retention", methods=["POST"])
@auth.role_required("manager")
def api_admin_org_retention(org_id: int):
    user = g.user
    u_role = auth.normalize_role(user["role"])
    data = request.json or {}
    days = int(data.get("retention_days", 30))
    
    if u_role != "admin" and org_id != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403
        
    database.set_org_retention(org_id, days)
    database.write_audit("org.retention_change", user_id=user["id"], org_id=org_id, resource_id=str(org_id), ip=_ip(), detail=str(days))
    return jsonify({"ok": True})


@app.route("/api/v1/admin/organizations/<int:org_id>/mfa-required", methods=["POST"])
@auth.role_required("manager")
def api_admin_org_mfa_required(org_id: int):
    user = g.user
    u_role = auth.normalize_role(user["role"])
    data = request.json or {}
    required = bool(data.get("mfa_required"))

    if u_role != "admin" and org_id != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403

    database.set_org_mfa_required(org_id, required)
    database.write_audit("org.mfa_required_change", user_id=user["id"], org_id=org_id, resource_id=str(org_id), ip=_ip(), detail=str(required))
    return jsonify({"ok": True})


# ── Admin API ──────────────────────────────────────────────────────────────────

@app.route("/api/v1/stats")
@auth.login_required
def api_stats():
    user = g.user
    if auth.normalize_role(user["role"]) == "admin":
        return jsonify(database.get_stats(org_id=None))
    return jsonify(database.get_stats(org_id=user["org_id"]))


@app.route("/api/v1/recent")
@auth.login_required
def api_recent():
    try:
        limit = min(int(request.args.get("limit", 10)), 50)
    except (TypeError, ValueError):
        limit = 10
    user = g.user
    if auth.normalize_role(user["role"]) == "admin":
        return jsonify(database.get_recent(limit, org_id=None))
    return jsonify(database.get_recent(limit, org_id=user["org_id"]))


@app.route("/api/v1/audit")
@auth.admin_required
def api_audit():
    try:
        limit = min(int(request.args.get("limit", 100)), 500)
    except (TypeError, ValueError):
        limit = 100
    return jsonify(database.get_audit_log(limit))


@app.route("/api/v1/citations")
@auth.login_required
def api_citations():
    """List all citations. Admins/reviewers can see drafts; normal users see verified only."""
    from detector.citation_db import _load
    db = _load()
    role = auth.normalize_role(g.user["role"])
    include_drafts = (role in ("admin", "reviewer"))

    out = []
    for fid, by_juris in db.items():
        for juris, rows in by_juris.items():
            for r in rows:
                if include_drafts or r.get("status") == "verified":
                    out.append({
                        "finding_id": fid,
                        "jurisdiction": juris,
                        "article": r.get("article"),
                        "source": r.get("source"),
                        "note": r.get("note"),
                        "status": r.get("status")
                    })
    return jsonify(out)


@app.route("/api/v1/citations/verify", methods=["POST"])
@auth.role_required("admin", "reviewer")
def api_verify_citation():
    """Transition a draft citation to verified status."""
    data = request.json or {}
    finding_id = data.get("finding_id")
    jurisdiction = data.get("jurisdiction")
    if not finding_id or not jurisdiction:
        return jsonify({"error": "Missing finding_id or jurisdiction"}), 400

    from detector.citation_db import verify_citation
    if verify_citation(finding_id, jurisdiction):
        database.write_audit(
            "cite.verify", user_id=g.user["id"], org_id=g.user["org_id"],
            resource_id=f"{finding_id}/{jurisdiction}", ip=_ip(),
        )
        return jsonify({"ok": True, "message": f"Citation {finding_id}/{jurisdiction} verified successfully"})
    else:
        return jsonify({"error": "Citation not found or status not changed"}), 404



# ── Tenant usage, history & professional review (Sprint 4) ─────────────────────

@app.route("/api/v1/usage", methods=["GET"])
@auth.role_required("manager")
def api_usage():
    usage = database.get_org_usage(g.user["org_id"])
    return jsonify(usage)


@app.route("/api/v1/history", methods=["GET"])
@auth.login_required
def api_history():
    user = g.user
    org_id = None if auth.normalize_role(user["role"]) == "admin" else user["org_id"]
    
    params = {
        "search": request.args.get("search"),
        "client": request.args.get("client"),
        "case_folder": request.args.get("case_folder"),
        "type": request.args.get("type"),
        "status": request.args.get("status"),
        "min_score": request.args.get("min_score"),
        "max_score": request.args.get("max_score"),
        "limit": request.args.get("limit"),
        "offset": request.args.get("offset"),
    }
    
    results = database.search_history(org_id, params)
    return jsonify(results)


@app.route("/api/v1/result/<analysis_id>/review", methods=["POST"])
@auth.role_required("reviewer")
def api_result_review(analysis_id: str):
    row = database.get_result(analysis_id)
    if row is None:
        return jsonify({"error": "Not found"}), 404
        
    user = g.user
    if auth.normalize_role(user["role"]) != "admin" and row.get("org_id") != user["org_id"]:
        return jsonify({"error": "Forbidden"}), 403
        
    data = request.json or {}
    status = data.get("status")
    comment = data.get("comment")
    
    valid_statuses = {"unreviewed", "confirmed", "edited", "rejected", "escalated"}
    if not status or status not in valid_statuses:
        return jsonify({"error": f"Invalid review status. Must be one of: {', '.join(valid_statuses)}"}), 400
        
    success = database.update_analysis_review(
        public_id=analysis_id,
        status=status,
        comment=comment,
        reviewer_email=user["email"]
    )
    if not success:
        return jsonify({"error": "Failed to update review status"}), 500
        
    database.write_audit(
        "analysis.review", user_id=user["id"], org_id=user["org_id"],
        resource_id=analysis_id, ip=_ip(), detail=f"status={status}"
    )
    return jsonify({"ok": True})



# ── PDF report ─────────────────────────────────────────────────────────────────

@app.route("/api/v1/report", methods=["POST"])
@auth.login_required
def report():
    # Check report subscription limit
    usage = database.get_org_usage(g.user["org_id"])
    if usage:
        report_limit = usage.get("report_limit", 50)
        report_used = usage.get("report_used", 0)
        if report_used + 1 > report_limit:
            return jsonify({"error": "Subscription limit exceeded: professional report allowance exhausted."}), 403

    from pdf_report import generate_pdf
    data = request.get_json(force=True, silent=True)
    if not data:
        return jsonify({"error": "Expected JSON body with analysis result"}), 400
    try:
        pdf_bytes = generate_pdf(data)
    except Exception as e:
        logger.exception("PDF generation failed")
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500

    # Increment report usage atomically only after successful PDF generation
    if usage:
        database.increment_org_usage(g.user["org_id"], reports=1)

    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": "attachment; filename=contract_risk_report.pdf"},
    )


# ── Legacy /analyze (kept for curl/API access) ────────────────────────────────

@app.route("/api/v1/analyze", methods=["POST"])
@auth.login_required
@limiter.limit("20 per minute")
def analyze():
    if auth.normalize_role(g.user["role"]) == "viewer":
        return jsonify({"error": "Forbidden: viewers cannot analyze documents"}), 403
    if os.getenv("LDV_PRODUCTION") == "1" and not crypto.is_enabled():
        return jsonify({"error": "Service configuration error: encryption is disabled or not configured in production"}), 500
    if os.getenv("LDV_ENFORCE_CONSENT", "1") == "1":
        if not (app.testing or app.config.get("TESTING") or os.environ.get("PYTEST_CURRENT_TEST") or os.environ.get("LDV_TESTING") == "1"):
            if not database.has_user_consented(g.user["id"]):
                return jsonify({
                    "error": "TOS and Privacy Policy acceptance required prior to document processing.",
                    "code": "CONSENT_REQUIRED",
                    "consent_url": "/api/v1/consent"
                }), 403
        
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    try:
        data, ext, text = _validate_and_extract(file)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

    try:
        lang = detect(text)
    except Exception:
        lang = "unknown"

    jurisdiction = detect_jurisdiction(text)
    policy_name = request.args.get("policy", "default_v1")
    result = _run_analysis(text, jurisdiction, lang, policy_name=policy_name)

    want_explain = request.args.get("explain", "0") == "1"
    if want_explain:
        layer1 = result["layer1"]
        layer2 = result["layer2"]
        layer3 = result["layer3"]
        analysis_text = text
        if lang not in ("en", "unknown"):
            try:
                analysis_text = translate_text(text, "en", src_lang=lang)
            except Exception:
                pass
        result["layer4"] = layer4_explain(
            analysis_text, jurisdiction=jurisdiction,
            layer1=layer1, layer2=layer2, layer3=layer3,
        )

    # ponytail: Gate raw_text behind ?debug=1 (P3 item 15)
    if request.args.get("debug") == "1":
        result["raw_text"] = text

    return jsonify(result)


# ── Health ─────────────────────────────────────────────────────────────────────

@app.route("/health")
@limiter.exempt
def health():
    from detector.detector_distilbert import is_available as l2_available
    from sydeco_engine import is_available as mlp_available
    
    db_ok = database.check_connection()
    
    # Check datasets
    datasets_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "datasets")
    required_csvs = [
        "abusive_clauses.csv", "dangerous_clauses_MASTERv2.csv",
        "illegal_clauses.csv", "leonine_clauses.csv",
        "required_clauses_MASTER.csv", "legal_citations.csv"
    ]
    datasets_ok = all(os.path.exists(os.path.join(datasets_dir, f)) for f in required_csvs)
    
    # Check model caches
    import hf_hub_connector
    qwen_cached = hf_hub_connector.is_model_cached("Qwen/Qwen3-1.7B")
    distilbert_cached = hf_hub_connector.is_model_cached("typeform/distilbert-base-uncased-mnli")
    
    try:
        from send_prompt import _model as qwen_model
        qwen_loaded = qwen_model is not None
    except Exception:
        qwen_loaded = False
        
    from translator_client import check_health as translator_health
    translator_status = translator_health()

    healthy = db_ok and datasets_ok
    status_str = "healthy" if healthy else "degraded"

    return jsonify({
        "status": status_str,
        "checks": {
            "database": "ready" if db_ok else "failed",
            "datasets": "ready" if datasets_ok else "missing",
            "model_cache": {
                "distilbert": "available" if distilbert_cached else "missing",
                "qwen3": "available" if qwen_cached else "missing"
            },
            # enabled=False when LDV_REMOTE_TRANSLATION isn't "local" or the
            # microservice isn't configured; doesn't affect overall "healthy"
            # status since remote translation is opt-in, not required.
            "lightml_translator": translator_status,
        },
        "layer1": "ready",
        "layer2_distilbert": l2_available(),
        "layer3_scorer": "ready",
        "layer4_qwen": qwen_loaded,
        "sydeco_mlp": mlp_available(),
        "encryption": {"enabled": crypto.is_enabled()},
        "retention_days": database.retention_days(),
    }), 200 if healthy else 500


# ── Frontend pages ─────────────────────────────────────────────────────────────

@app.route("/")
def home():
    return send_from_directory(FRONTEND_DIR, "index.html")


@app.route("/result")
@app.route("/result/<analysis_id>")
def result_page(analysis_id=None):
    return send_from_directory(FRONTEND_DIR, "result.html")


@app.route("/admin")
def admin_page():
    user = auth.current_user()
    if user is None:
        return redirect("/login")
    return send_from_directory(FRONTEND_DIR, "admin.html")


@app.route("/citations")
def citation_review_page():
    user = auth.current_user()
    if user is None or auth.normalize_role(user["role"]) not in ("admin", "reviewer"):
        return redirect("/login")
    return send_from_directory(FRONTEND_DIR, "citations.html")


@app.route("/account")
def account_page():
    user = auth.current_user()
    if user is None:
        return redirect("/login")
    return send_from_directory(FRONTEND_DIR, "account.html")


@app.route("/swagger.json")
def swagger_json():
    return send_from_directory(FRONTEND_DIR, "swagger.json")


@app.route("/docs")
@app.route("/swagger")
def swagger_docs():
    return send_from_directory(FRONTEND_DIR, "swagger.html")


@app.route("/<path:filename>")
def frontend_files(filename):
    filepath = os.path.join(FRONTEND_DIR, filename)
    if os.path.isfile(filepath):
        return send_from_directory(FRONTEND_DIR, filename)
    return send_from_directory(FRONTEND_DIR, "index.html")


# ── Profile Admin APIs ──────────────────────────────────────────────────────────

@app.route("/api/v1/admin/profiles", methods=["GET"])
@auth.admin_required
def admin_list_profiles():
    """List all profiles in the registry (both active and inactive)."""
    try:
        from detector.detector_profiles import ProfileManager
        manager = ProfileManager()
        
        # Load active and inactive profile lists
        active_ids = manager._registry.get("profiles", {}).keys()
        inactive_ids = manager._registry.get("inactive_profiles", {}).keys()
        
        profiles = []
        for pid in list(active_ids) + list(inactive_ids):
            try:
                p = manager.get_profile(pid)
                metadata = p.get("metadata", {})
                profiles.append({
                    "profile_id": pid,
                    "version": p.get("version"),
                    "validation_status": p.get("validation_status"),
                    "review_date": p.get("review_date"),
                    "display_name": metadata.get("display_name"),
                    "family": metadata.get("family"),
                    "active": pid in active_ids
                })
            except Exception as err:
                profiles.append({
                    "profile_id": pid,
                    "active": pid in active_ids,
                    "error": str(err)
                })
                
        return jsonify({"profiles": profiles})
    except Exception as exc:
        return jsonify({"error": f"Failed to list profiles: {exc}"}), 500


@app.route("/api/v1/admin/profiles/<profile_id>", methods=["GET"])
@auth.admin_required
def admin_get_profile(profile_id):
    """Retrieve details of a specific profile by ID."""
    try:
        from detector.detector_profiles import ProfileManager
        manager = ProfileManager()
        p = manager.get_profile(profile_id)
        active = profile_id in manager._registry.get("profiles", {})
        return jsonify({
            "profile": p,
            "active": active
        })
    except ValueError:
        try:
            from detector.detector_profiles import ProfileManager
            manager = ProfileManager()
            inactive_map = manager._registry.get("inactive_profiles", {})
            if profile_id in inactive_map:
                filename = inactive_map[profile_id]
                profile_path = os.path.join(os.path.dirname(manager.registry_path), filename)
                with open(profile_path, "r", encoding="utf-8") as f:
                    p = json.load(f)
                return jsonify({
                    "profile": p,
                    "active": False
                })
        except Exception:
            pass
        return jsonify({"error": f"Profile '{profile_id}' not found."}), 404
    except Exception as exc:
        return jsonify({"error": f"Failed to retrieve profile: {exc}"}), 500


@app.route("/api/v1/admin/profiles/validate", methods=["POST"])
@auth.admin_required
def admin_validate_profile_payload():
    """Validate a raw profile payload against JSON schema and rules."""
    payload = request.get_json()
    if not payload:
        return jsonify({"error": "No JSON payload provided"}), 400
        
    try:
        from detector.detector_profiles import ProfileManager, validate_profile
        manager = ProfileManager()
        validate_profile(payload, manager._rules)
        return jsonify({"status": "valid", "message": "Profile payload conforms to schema and rules classification."})
    except Exception as exc:
        return jsonify({"status": "invalid", "error": str(exc)}), 400


@app.route("/api/v1/admin/profiles/<profile_id>/publish", methods=["POST"])
@auth.admin_required
def admin_publish_profile(profile_id):
    """Publish or update a profile specification payload."""
    payload = request.get_json()
    if not payload:
        return jsonify({"error": "No JSON payload provided"}), 400
        
    try:
        from detector.detector_profiles import ProfileManager
        manager = ProfileManager()
        manager.publish_profile(profile_id, payload)
        return jsonify({"status": "success", "message": f"Profile '{profile_id}' published successfully and activated."})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@app.route("/api/v1/admin/profiles/<profile_id>/activate", methods=["POST"])
@auth.admin_required
def admin_activate_profile(profile_id):
    """Activate a deactivated profile."""
    try:
        from detector.detector_profiles import ProfileManager
        manager = ProfileManager()
        manager.activate_profile(profile_id)
        return jsonify({"status": "success", "message": f"Profile '{profile_id}' activated successfully."})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@app.route("/api/v1/admin/profiles/<profile_id>/deactivate", methods=["POST"])
@auth.admin_required
def admin_deactivate_profile(profile_id):
    """Deactivate an active profile."""
    try:
        from detector.detector_profiles import ProfileManager
        manager = ProfileManager()
        manager.deactivate_profile(profile_id)
        return jsonify({"status": "success", "message": f"Profile '{profile_id}' deactivated successfully."})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


@app.route("/api/v1/admin/profiles/<profile_id>/rollback", methods=["POST"])
@auth.admin_required
def admin_rollback_profile(profile_id):
    """Roll back a profile to a previous version stored in history."""
    params = request.get_json() or {}
    target_version = params.get("version")
    if not target_version:
        return jsonify({"error": "Missing 'version' parameter in payload"}), 400
        
    try:
        from detector.detector_profiles import ProfileManager
        manager = ProfileManager()
        manager.rollback_profile(profile_id, target_version)
        return jsonify({"status": "success", "message": f"Profile '{profile_id}' successfully rolled back to version {target_version}."})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400


if __name__ == "__main__":
    # Debug mode exposes the Werkzeug debugger (remote code execution if the
    # port is reachable) — opt-in only. Production: gunicorn -w 2 app:app
    app.run(debug=os.getenv("LDV_DEBUG", "0") == "1")
